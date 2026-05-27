# -*- coding: utf-8 -*-
"""
Koib-V-4.6 — Модуль парсинга документов
★ ИСПРАВЛЕНО: явный вызов gc.collect() после каждой страницы — защита от OOM
★ ИСПРАВЛЕНО: явное удаление pixmap и изображений
"""
import io
import re
import gc
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field, asdict

import fitz
from docx import Document as DocxDocument
from PIL import Image

from .utils import (
    clean_text, text_hash, detect_model_in_text,
    detect_model_from_filename, find_figure_caption,
    extract_headings, estimate_tokens, generate_unique_id,
)
from config import OCR_DPI, OCR_MIN_TEXT_CHARS, MIN_IMAGE_WIDTH, MIN_IMAGE_HEIGHT, FIGURES_DIR

logger = logging.getLogger("koib.parsing")


@dataclass
class DocumentElement:
    content: str
    element_type: str
    source: str = ""
    page: int = 0
    heading: str = ""
    model: str = "unknown"
    element_id: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __post_init__(self):
        if not self.element_id:
            self.element_id = text_hash(
                f"{self.source}:{self.page}:{self.element_type}:{self.content[:200]}"
            )

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)

    @property
    def is_structured(self) -> bool:
        return self.element_type in ("table", "formula", "figure")


def _expand_rect(rect: fitz.Rect, margin: float) -> fitz.Rect:
    return fitz.Rect(rect.x0 - margin, rect.y0 - margin,
                     rect.x1 + margin, rect.y1 + margin)


def _is_scanned_page(page: fitz.Page, min_chars: int = OCR_MIN_TEXT_CHARS) -> bool:
    text = page.get_text("text").strip()
    if len(text) >= min_chars:
        return False
    images = page.get_images(full=True)
    if not images:
        return len(text) < min_chars
    page_area = page.rect.width * page.rect.height
    for img_info in images:
        try:
            xref = img_info[0]
            base_image = page.parent.extract_image(xref)
            if not base_image:
                continue
            img = Image.open(io.BytesIO(base_image["image"]))
            if img.width * img.height / page_area > 0.8:
                img.close()
                return True
            img.close()
        except Exception:
            continue
    return True


def _ocr_image(image_pil: Image.Image, lang: str = "rus+eng") -> str:
    if image_pil is None:
        return ""
    try:
        import pytesseract
        text = clean_text(pytesseract.image_to_string(
            image_pil, lang=lang, config="--psm 6"
        ))
        if len(text) >= 30:
            return text
    except Exception as exc:
        logger.debug(f"Tesseract OCR error: {exc}")
    return ""


def _extract_tables_from_page(page: fitz.Page) -> List[Dict[str, Any]]:
    tables = []
    try:
        tab_finder = page.find_tables()
        for tab in tab_finder:
            try:
                rows = tab.extract()
                if not rows or len(rows) < 2:
                    continue
                md_lines = []
                num_cols = max(len(r) for r in rows) if rows else 0
                for i, row in enumerate(rows):
                    cells = [str(c).strip() if c else "" for c in row]
                    while len(cells) < num_cols:
                        cells.append("")
                    md_lines.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        md_lines.append("| " + " | ".join(["---"] * num_cols) + " |")
                markdown = "\n".join(md_lines)
                tables.append({
                    "text": markdown,
                    "num_rows": len(rows),
                    "num_cols": num_cols,
                    "bbox": tuple(tab.bbox) if hasattr(tab, "bbox") else (0, 0, 0, 0),
                })
            except Exception as exc:
                logger.debug(f"Ошибка конвертации таблицы: {exc}")
    except AttributeError:
        logger.debug("PyMuPDF < 1.23: find_tables() недоступен")
    except Exception as exc:
        logger.debug(f"Ошибка поиска таблиц: {exc}")
    return tables


def _detect_formulas_in_text(text: str) -> List[Dict[str, Any]]:
    formulas = []
    for match in re.finditer(r'\$([^$]+)\$', text):
        formulas.append({
            "content": match.group(1).strip(),
            "formula_type": "latex_inline",
            "start": match.start(), "end": match.end(),
        })
    for match in re.finditer(r'\$\$(.+?)\$\$', text, re.DOTALL):
        formulas.append({
            "content": match.group(1).strip(),
            "formula_type": "latex_block",
            "start": match.start(), "end": match.end(),
        })
    math_pattern = re.compile(r'[=+\-*/^√∑∫∏∂∇∞≈≠≤≥±αβγδεζηθλμπρσφψω]')
    for line in text.split('\n'):
        line = line.strip()
        if len(line) < 5:
            continue
        if math_pattern.search(line) and not line.startswith('|'):
            already_found = any(
                f["start"] <= text.find(line) <= f["end"] for f in formulas
            )
            if not already_found:
                formulas.append({
                    "content": line,
                    "formula_type": "suspected_formula",
                    "start": text.find(line),
                    "end": text.find(line) + len(line),
                })
    return formulas


def parse_pdf(file_path: Path, model_hint: str = "") -> List[DocumentElement]:
    if not file_path.exists():
        logger.error(f"Файл не найден: {file_path}")
        return []
    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    elements: List[DocumentElement] = []

    try:
        doc = fitz.open(str(file_path))
    except Exception as exc:
        logger.error(f"Не удалось открыть PDF {filename}: {exc}")
        return []

    logger.info(f"Парсинг PDF: {filename} ({len(doc)} стр.)")

    full_text_sample = ""
    for page in doc:
        full_text_sample += page.get_text("text") + "\n"
    detected_model, confidence = detect_model_in_text(full_text_sample)
    if confidence > 0.3:
        model = detected_model

    current_heading = ""
    for page_num in range(len(doc)):
        page = doc[page_num]
        page_text = page.get_text("text").strip()

        # ★ ИСПРАВЛЕНО: OCR с контролируемым жизненным циклом объектов
        if _is_scanned_page(page):
            pix = page.get_pixmap(dpi=OCR_DPI)
            try:
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = _ocr_image(img)
                if ocr_text:
                    elements.append(DocumentElement(
                        content=ocr_text, element_type="text",
                        source=filename, page=page_num + 1, model=model,
                        heading=current_heading,
                        metadata={"ocr": True},
                    ))
                img.close()
            finally:
                # ★ КРИТИЧНО: немедленное освобождение битмапа (~15-25 МБ)
                del pix
                gc.collect()
            continue

        headings = extract_headings(page_text)
        if headings:
            current_heading = headings[0]

        tables = _extract_tables_from_page(page)
        for table_data in tables:
            table_text = clean_text(table_data["text"])
            if table_text:
                detected_model, conf = detect_model_in_text(table_text)
                elements.append(DocumentElement(
                    content=table_text, element_type="table",
                    source=filename, page=page_num + 1,
                    model=detected_model if conf > 0.3 else model,
                    heading=current_heading,
                    metadata={
                        "num_rows": table_data["num_rows"],
                        "num_cols": table_data["num_cols"],
                        "bbox": table_data["bbox"],
                    },
                ))

        formulas = _detect_formulas_in_text(page_text)
        for formula_data in formulas:
            elements.append(DocumentElement(
                content=formula_data["content"], element_type="formula",
                source=filename, page=page_num + 1, model=model,
                heading=current_heading,
                metadata={"formula_type": formula_data["formula_type"]},
            ))

        images = page.get_images(full=True)
        for img_idx, img_info in enumerate(images):
            img = None
            try:
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue
                img = Image.open(io.BytesIO(base_image["image"]))
                if img.width < MIN_IMAGE_WIDTH or img.height < MIN_IMAGE_HEIGHT:
                    continue
                caption = find_figure_caption(page_text)
                content = caption if caption else f"Изображение {img_idx + 1}"
                FIGURES_DIR.mkdir(parents=True, exist_ok=True)
                img_filename = f"{file_path.stem}_p{page_num + 1}_img{img_idx}.png"
                img_path = FIGURES_DIR / img_filename
                img.save(str(img_path))
                elements.append(DocumentElement(
                    content=content, element_type="figure",
                    source=filename, page=page_num + 1, model=model,
                    heading=current_heading,
                    metadata={
                        "image_path": str(img_path),
                        "width": img.width, "height": img.height,
                    },
                ))
            except Exception as exc:
                logger.debug(f"Ошибка изображения: {exc}")
            finally:
                if img is not None:
                    try:
                        img.close()
                    except Exception:
                        pass

        if page_text:
            for heading in headings:
                elements.append(DocumentElement(
                    content=heading, element_type="heading",
                    source=filename, page=page_num + 1, model=model,
                    heading=heading,
                ))
            cleaned = clean_text(page_text)
            if len(cleaned) >= OCR_MIN_TEXT_CHARS:
                detected_model, conf = detect_model_in_text(cleaned)
                elements.append(DocumentElement(
                    content=cleaned, element_type="text",
                    source=filename, page=page_num + 1,
                    model=detected_model if conf > 0.3 else model,
                    heading=current_heading,
                ))

        # ★ ИСПРАВЛЕНО: принудительная сборка мусора после каждой страницы
        gc.collect()

    doc.close()
    # ★ Финальная очистка после всего документа
    gc.collect()
    logger.info(f"Извлечено {len(elements)} элементов из {filename}")
    return elements


def parse_docx(file_path: Path, model_hint: str = "") -> List[DocumentElement]:
    if not file_path.exists():
        logger.error(f"Файл не найден: {file_path}")
        return []
    filename = file_path.name
    model = model_hint or detect_model_from_filename(filename)
    elements: List[DocumentElement] = []

    try:
        doc = DocxDocument(str(file_path))
    except Exception as exc:
        logger.error(f"Не удалось открыть DOCX {filename}: {exc}")
        return []

    logger.info(f"Парсинг DOCX: {filename}")
    current_heading = ""

    for table_idx, table in enumerate(doc.tables):
        rows_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not rows_data:
            continue
        num_cols = max(len(r) for r in rows_data)
        header = rows_data[0] if rows_data else []
        while len(header) < num_cols:
            header.append(f"Кол.{len(header) + 1}")
        md_lines = ["| " + " | ".join(header) + " |",
                    "| " + " | ".join(["---"] * num_cols) + " |"]
        for row in rows_data[1:]:
            while len(row) < num_cols:
                row.append("")
            md_lines.append("| " + " | ".join(row[:num_cols]) + " |")
        table_md = "\n".join(md_lines)
        detected_model, conf = detect_model_in_text(table_md)
        elements.append(DocumentElement(
            content=clean_text(table_md), element_type="table",
            source=filename, page=0,
            model=detected_model if conf > 0.3 else model,
            heading=current_heading,
            metadata={
                "num_rows": len(rows_data), "num_cols": num_cols,
                "table_index": table_idx,
            },
        ))

    text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name.lower() if para.style else ""
        is_heading = "heading" in style_name or "заголовок" in style_name
        if is_heading:
            current_heading = text
            detected_model, conf = detect_model_in_text(text)
            elements.append(DocumentElement(
                content=clean_text(text), element_type="heading",
                source=filename, page=0,
                model=detected_model if conf > 0.3 else model,
                heading=text,
            ))
        else:
            text_parts.append(text)

    if text_parts:
        combined = "\n".join(text_parts)
        cleaned = clean_text(combined)
        if len(cleaned) >= 50:
            detected_model, conf = detect_model_in_text(cleaned)
            elements.append(DocumentElement(
                content=cleaned, element_type="text",
                source=filename, page=0,
                model=detected_model if conf > 0.3 else model,
                heading=current_heading,
            ))

    gc.collect()
    logger.info(f"Извлечено {len(elements)} элементов из {filename}")
    return elements


def parse_document(file_path: Path, engine: Optional[str] = None,
                   model_hint: str = "") -> List[DocumentElement]:
    file_path = Path(file_path)
    if file_path.suffix.lower() == '.pdf':
        return parse_pdf(file_path, model_hint)
    elif file_path.suffix.lower() in ('.docx', '.doc'):
        return parse_docx(file_path, model_hint)
    return []